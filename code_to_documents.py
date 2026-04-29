#!/usr/bin/env python3
"""
Code Files to Documents Converter v2.0
Reads all code files from a project directory, consolidates them into one folder,
and splits them into multiple documents.

New in v2.0:
  - Folder exclusions (node_modules, .git, build, etc.)
  - Config file save/load (no more re-entering 6 prompts)
  - CLI arguments via argparse (fully scriptable)
  - Index document generation
  - Dry-run mode
  - Progress bar (requires tqdm, optional)
  - HTML output with syntax highlighting (requires pygments, optional)
  - Table of contents in each document
  - Function/class extractor per file
  - Stats summary per document
"""

import os
import sys
import shutil
import json
import argparse
import re
import datetime
from pathlib import Path

# ─── Optional dependencies ────────────────────────────────────────────────────
try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False

try:
    from pygments import highlight
    from pygments.lexers import get_lexer_for_filename, TextLexer
    from pygments.formatters import HtmlFormatter
    from pygments.util import ClassNotFound
    HAS_PYGMENTS = True
except ImportError:
    HAS_PYGMENTS = False

# ─── Constants ────────────────────────────────────────────────────────────────
CONFIG_FILE = "code_to_docs_config.json"

DEFAULT_EXCLUDES = {
    "node_modules", ".git", "__pycache__", "build", "dist",
    ".venv", "venv", "env", "target", ".idea", ".vscode",
    "out", "bin", "obj", ".next", ".nuxt", "coverage",
    ".gradle", ".mvn", "vendor", "bower_components",
}

COMMON_EXTENSIONS = {
    "Java/Kotlin/Scala": [".java", ".kt", ".scala"],
    "JavaScript/TypeScript": [".js", ".jsx", ".ts", ".tsx", ".mjs"],
    "Python": [".py"],
    "C/C++": [".c", ".cpp", ".h", ".hpp"],
    "C#": [".cs"],
    "Go": [".go"],
    "Rust": [".rs"],
    "PHP": [".php"],
    "Ruby": [".rb"],
    "Swift/ObjC": [".swift", ".m", ".mm"],
    "Web": [".html", ".htm", ".css", ".scss", ".sass"],
    "Data/Config": [".json", ".xml", ".yaml", ".yml"],
    "Docs": [".md", ".txt"],
}

# Regex patterns to extract class/function names per language
EXTRACTION_PATTERNS = {
    ".java":  [r"(?:public|private|protected|static|\s)+(?:class|interface|enum)\s+(\w+)",
               r"(?:public|private|protected|static|\s)+\w+\s+(\w+)\s*\("],
    ".kt":    [r"(?:class|object|interface|fun)\s+(\w+)"],
    ".py":    [r"^(?:class)\s+(\w+)", r"^(?:def)\s+(\w+)"],
    ".js":    [r"(?:function\s+(\w+)|class\s+(\w+)|const\s+(\w+)\s*=\s*(?:async\s*)?\()"],
    ".jsx":   [r"(?:function\s+(\w+)|class\s+(\w+)|const\s+(\w+)\s*=\s*(?:async\s*)?\()"],
    ".ts":    [r"(?:function\s+(\w+)|class\s+(\w+)|interface\s+(\w+)|const\s+(\w+)\s*=)"],
    ".tsx":   [r"(?:function\s+(\w+)|class\s+(\w+)|interface\s+(\w+)|const\s+(\w+)\s*=)"],
    ".cs":    [r"(?:class|interface|struct|enum)\s+(\w+)", r"(?:public|private|protected|static|\s)+\w+\s+(\w+)\s*\("],
    ".go":    [r"^func\s+(?:\(\w+\s+\*?\w+\)\s+)?(\w+)", r"^type\s+(\w+)\s+(?:struct|interface)"],
    ".rs":    [r"(?:fn|struct|enum|trait|impl)\s+(\w+)"],
    ".php":   [r"(?:function|class)\s+(\w+)"],
    ".rb":    [r"(?:def|class|module)\s+(\w+)"],
    ".swift": [r"(?:func|class|struct|enum|protocol)\s+(\w+)"],
    ".cpp":   [r"(?:class|struct)\s+(\w+)", r"\b(\w+)\s*\([^)]*\)\s*(?:const\s*)?\{"],
    ".c":     [r"\b(\w+)\s*\([^)]*\)\s*\{"],
    ".h":     [r"(?:class|struct)\s+(\w+)", r"\b(\w+)\s*\("],
}


# ─── Config ───────────────────────────────────────────────────────────────────

def load_config():
    """Load saved config from disk if it exists."""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_config(config: dict):
    """Save current run config to disk for future reuse."""
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=2)
        print(f"  Config saved to: {CONFIG_FILE}")
    except Exception as e:
        print(f"  (Could not save config: {e})")


# ─── File Discovery ────────────────────────────────────────────────────────────

def get_all_files_by_extension(search_dir, extensions, exclude_dirs=None):
    """
    Recursively find all files matching given extensions, skipping excluded folders.
    """
    if exclude_dirs is None:
        exclude_dirs = DEFAULT_EXCLUDES

    normalized_exts = []
    for ext in extensions:
        ext = ext.strip()
        if not ext.startswith("."):
            ext = "." + ext
        normalized_exts.append(ext.lower())

    files = []
    for root, dirs, filenames in os.walk(search_dir):
        # Skip excluded directories in-place so os.walk won't descend into them
        dirs[:] = [d for d in dirs if d not in exclude_dirs and not d.startswith(".")]
        for filename in filenames:
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in normalized_exts:
                files.append(os.path.join(root, filename))

    return sorted(files)


def count_files_preview(project_dir, extensions, exclude_dirs):
    """
    Dry-run preview: count matching files and show breakdown by extension,
    and report how many files were skipped in excluded folders.
    """
    # Count included files
    included = get_all_files_by_extension(project_dir, extensions, exclude_dirs)

    # Count skipped files (walk without exclusions, count what's excluded)
    normalized_exts = []
    for ext in extensions:
        ext = ext.strip()
        if not ext.startswith("."):
            ext = "." + ext
        normalized_exts.append(ext.lower())

    skipped = []
    for root, dirs, filenames in os.walk(project_dir):
        folder_name = os.path.basename(root)
        if folder_name in exclude_dirs or folder_name.startswith("."):
            for filename in filenames:
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext in normalized_exts:
                    skipped.append(os.path.join(root, filename))

    # Breakdown by extension
    ext_counts = {}
    for f in included:
        ext = os.path.splitext(f)[1].lower()
        ext_counts[ext] = ext_counts.get(ext, 0) + 1

    print("\n" + "─" * 60)
    print("DRY RUN — No files will be written")
    print("─" * 60)
    print(f"\n  Would process : {len(included)} files")
    print(f"  Would skip    : {len(skipped)} files (in excluded folders)\n")

    if ext_counts:
        print("  Breakdown by extension:")
        for ext, count in sorted(ext_counts.items()):
            print(f"    {ext:<10} {count} files")

    if included:
        print(f"\n  First 5 files that would be included:")
        for f in included[:5]:
            print(f"    {f}")
        if len(included) > 5:
            print(f"    ... and {len(included) - 5} more")

    print("\n" + "─" * 60)
    print("Run without --dry-run to process these files.")
    print("─" * 60 + "\n")


# ─── Code Intelligence ─────────────────────────────────────────────────────────

def extract_symbols(file_path, content):
    """Extract class/function names from a file using regex patterns."""
    ext = os.path.splitext(file_path)[1].lower()
    patterns = EXTRACTION_PATTERNS.get(ext, [])
    if not patterns:
        return [], []

    classes = []
    functions = []
    lines = content.split("\n")

    for line in lines:
        for pattern in patterns:
            for match in re.finditer(pattern, line):
                groups = [g for g in match.groups() if g]
                for name in groups:
                    if len(name) > 1 and name not in ("if", "for", "while", "return",
                                                       "new", "try", "catch", "import",
                                                       "from", "int", "void", "bool",
                                                       "string", "var", "let", "const"):
                        if re.match(r"^[A-Z]", name):
                            if name not in classes:
                                classes.append(name)
                        else:
                            if name not in functions:
                                functions.append(name)

    return classes[:10], functions[:15]  # Cap to keep things readable


def extract_todos(content):
    """Find TODO/FIXME/HACK/NOTE comments in file content."""
    todos = []
    for i, line in enumerate(content.split("\n"), 1):
        if re.search(r"\b(TODO|FIXME|HACK|XXX|NOTE)\b", line, re.IGNORECASE):
            stripped = line.strip()
            if len(stripped) > 80:
                stripped = stripped[:77] + "..."
            todos.append((i, stripped))
    return todos[:5]  # Cap at 5 per file


def file_stats(content):
    """Return line count breakdown for a file."""
    lines = content.split("\n")
    total = len(lines)
    blank = sum(1 for l in lines if not l.strip())
    comment = sum(1 for l in lines if re.match(r"^\s*(//|#|/\*|\*|<!--)", l.strip()))
    code = total - blank - comment
    return {"total": total, "code": code, "blank": blank, "comment": comment}


# ─── File Consolidation ────────────────────────────────────────────────────────

def consolidate_files(project_dir, consolidated_dir, extensions, exclude_dirs, dry_run=False):
    """Copy all matching files into a single flat folder, handling duplicates."""
    print(f"\n[1/3] Consolidating files...")

    os.makedirs(consolidated_dir, exist_ok=True)
    files = get_all_files_by_extension(project_dir, extensions, exclude_dirs)

    if not files:
        print(f"  No files found with extensions: {extensions}")
        return []

    print(f"  Found {len(files)} files")

    filename_counts = {}
    consolidated_files = []

    iterator = tqdm(files, desc="  Copying") if HAS_TQDM else files

    for original_path in iterator:
        filename = os.path.basename(original_path)

        if filename in filename_counts:
            filename_counts[filename] += 1
            rel_path = os.path.relpath(original_path, project_dir)
            path_parts = rel_path.split(os.sep)
            if len(path_parts) > 1:
                new_filename = f"{path_parts[-2]}_{filename}"
            else:
                new_filename = f"file_{filename_counts[filename]}_{filename}"
        else:
            filename_counts[filename] = 0
            new_filename = filename

        dest_path = os.path.join(consolidated_dir, new_filename)

        counter = 1
        base, ext = os.path.splitext(new_filename)
        while os.path.exists(dest_path):
            dest_path = os.path.join(consolidated_dir, f"{base}_{counter}{ext}")
            counter += 1

        if not dry_run:
            shutil.copy2(original_path, dest_path)
        consolidated_files.append(dest_path)

    print(f"  ✓ Consolidated {len(consolidated_files)} files → {consolidated_dir}")
    return consolidated_files


# ─── Output: Text ─────────────────────────────────────────────────────────────

def format_file_block_text(file_path, content, original_path=None):
    """Format a single file for text output, with symbol summary and TODOs."""
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].upper().replace(".", "")
    classes, functions = extract_symbols(file_path, content)
    todos = extract_todos(content)
    stats = file_stats(content)
    display_path = original_path or file_path

    out = f"{'='*80}\n"
    out += f"FILE: {filename}\n"
    out += f"TYPE: {file_ext}   |   Lines: {stats['total']} total / {stats['code']} code / {stats['blank']} blank / {stats['comment']} comment\n"
    out += f"PATH: {display_path}\n"

    if classes:
        out += f"CLASSES: {', '.join(classes)}\n"
    if functions:
        out += f"FUNCTIONS: {', '.join(functions)}\n"
    if todos:
        out += f"TODOs ({len(todos)}):\n"
        for lineno, text in todos:
            out += f"  Line {lineno}: {text}\n"

    out += f"{'='*80}\n\n"
    out += content
    out += f"\n\n{'='*80}\n"
    out += f"END OF FILE: {filename}\n"
    out += f"{'='*80}\n\n"
    return out, stats


def create_text_documents(consolidated_dir, output_dir, file_exts, num_docs=10):
    """Create .txt output documents with TOC, symbol summaries, and stats."""
    print(f"\n[2/3] Creating text documents...")

    files = get_all_files_by_extension(consolidated_dir, file_exts)
    total_files = len(files)

    if total_files == 0:
        print("  No files found!")
        return

    print(f"  Processing {total_files} files across {num_docs} documents")
    os.makedirs(output_dir, exist_ok=True)

    ext_str = "_".join([ext.replace(".", "") for ext in file_exts])
    doc_groups = split_into_documents(files, num_docs)

    for i, group in enumerate(doc_groups, 1):
        output_file = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.txt")
        doc_stats = {"total_lines": 0, "code_lines": 0, "blank_lines": 0, "comment_lines": 0}

        with open(output_file, "w", encoding="utf-8") as out:
            # Header
            out.write(f"{'#'*80}\n")
            out.write(f"Code Files Collection — Part {i} of {num_docs}\n")
            out.write(f"File Types : {', '.join(file_exts)}\n")
            out.write(f"Files      : {len(group)}\n")
            out.write(f"Generated  : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            out.write(f"{'#'*80}\n\n")

            # Table of Contents
            out.write("TABLE OF CONTENTS\n")
            out.write("─" * 40 + "\n")
            for j, fp in enumerate(group, 1):
                out.write(f"  {j:>3}. {os.path.basename(fp)}\n")
            out.write("\n" + "─" * 40 + "\n\n")

            # File contents
            iterator = tqdm(group, desc=f"  Part {i:02d}") if HAS_TQDM else group
            for file_path in iterator:
                content = read_file_content(file_path)
                if content:
                    formatted, stats = format_file_block_text(file_path, content)
                    out.write(formatted)
                    for k in doc_stats:
                        key = k.replace("_lines", "")
                        doc_stats[k] += stats.get(key if key != "total" else "total", 0)

            # Document stats footer
            out.write(f"\n{'#'*80}\n")
            out.write(f"DOCUMENT STATS\n")
            out.write(f"  Total lines  : {doc_stats['total_lines']:,}\n")
            out.write(f"  Code lines   : {doc_stats['code_lines']:,}\n")
            out.write(f"  Blank lines  : {doc_stats['blank_lines']:,}\n")
            out.write(f"  Comment lines: {doc_stats['comment_lines']:,}\n")
            out.write(f"{'#'*80}\n")

        print(f"  ✓ {output_file}  ({len(group)} files)")

    print(f"\n  Created {num_docs} text documents in: {output_dir}")


# ─── Output: HTML ─────────────────────────────────────────────────────────────

def create_html_documents(consolidated_dir, output_dir, file_exts, num_docs=10):
    """Create syntax-highlighted HTML output documents."""
    if not HAS_PYGMENTS:
        print("\n  pygments not installed — falling back to text output.")
        print("  Install with: pip install pygments\n")
        create_text_documents(consolidated_dir, output_dir, file_exts, num_docs)
        return

    print(f"\n[2/3] Creating HTML documents (syntax highlighted)...")

    files = get_all_files_by_extension(consolidated_dir, file_exts)
    total_files = len(files)

    if total_files == 0:
        print("  No files found!")
        return

    print(f"  Processing {total_files} files across {num_docs} documents")
    os.makedirs(output_dir, exist_ok=True)

    formatter = HtmlFormatter(style="monokai", linenos=True, cssclass="code")
    css = formatter.get_style_defs(".code")
    ext_str = "_".join([ext.replace(".", "") for ext in file_exts])
    doc_groups = split_into_documents(files, num_docs)

    for i, group in enumerate(doc_groups, 1):
        output_file = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.html")

        with open(output_file, "w", encoding="utf-8") as out:
            # Build TOC
            toc_items = ""
            for j, fp in enumerate(group, 1):
                fn = os.path.basename(fp)
                anchor = f"file-{j}"
                toc_items += f'<li><a href="#{anchor}">{fn}</a></li>\n'

            out.write(f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Code Collection — Part {i} of {num_docs}</title>
<style>
  body {{ font-family: Arial, sans-serif; background: #1e1e1e; color: #d4d4d4; margin: 0; padding: 0; }}
  .sidebar {{ position: fixed; top: 0; left: 0; width: 260px; height: 100vh; background: #252526;
              overflow-y: auto; padding: 16px; box-sizing: border-box; border-right: 1px solid #444; }}
  .sidebar h2 {{ color: #9cdcfe; font-size: 13px; text-transform: uppercase; margin: 0 0 12px; }}
  .sidebar ol {{ padding-left: 18px; margin: 0; }}
  .sidebar li {{ margin-bottom: 4px; }}
  .sidebar a {{ color: #ce9178; text-decoration: none; font-size: 12px; word-break: break-all; }}
  .sidebar a:hover {{ color: #fff; }}
  .content {{ margin-left: 280px; padding: 24px; max-width: 1100px; }}
  .file-block {{ margin-bottom: 40px; border: 1px solid #333; border-radius: 6px; overflow: hidden; }}
  .file-header {{ background: #2d2d2d; padding: 10px 16px; border-bottom: 1px solid #444; }}
  .file-header h3 {{ margin: 0; color: #9cdcfe; font-size: 14px; }}
  .file-meta {{ font-size: 12px; color: #888; margin-top: 4px; }}
  .file-meta span {{ margin-right: 16px; }}
  .symbols {{ font-size: 12px; color: #b5cea8; margin-top: 4px; }}
  .todos {{ font-size: 12px; color: #f48771; margin-top: 4px; }}
  .stats-footer {{ background: #252526; padding: 12px 16px; margin-top: 40px;
                   border-radius: 6px; font-size: 13px; color: #888; }}
  {css}
  .highlight {{ margin: 0; border-radius: 0; }}
</style>
</head>
<body>
<div class="sidebar">
  <h2>Part {i} of {num_docs}</h2>
  <p style="color:#888;font-size:11px;margin:0 0 12px">{len(group)} files &bull; {', '.join(file_exts)}</p>
  <ol>{toc_items}</ol>
</div>
<div class="content">
  <h1 style="color:#9cdcfe;font-size:18px">Code Collection — Part {i} of {num_docs}</h1>
  <p style="color:#888;font-size:13px">Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
""")
            total_lines = 0
            for j, file_path in enumerate(group, 1):
                content = read_file_content(file_path)
                if not content:
                    continue

                filename = os.path.basename(file_path)
                ext = os.path.splitext(filename)[1].lower()
                classes, functions = extract_symbols(file_path, content)
                todos = extract_todos(content)
                stats = file_stats(content)
                total_lines += stats["total"]

                try:
                    lexer = get_lexer_for_filename(filename)
                except ClassNotFound:
                    lexer = TextLexer()

                highlighted = highlight(content, lexer, formatter)

                classes_html = f'<div class="symbols">Classes: {", ".join(classes)}</div>' if classes else ""
                funcs_html = f'<div class="symbols">Functions: {", ".join(functions)}</div>' if functions else ""
                todos_html = ""
                if todos:
                    todo_list = " &bull; ".join([f"L{ln}: {t}" for ln, t in todos[:3]])
                    todos_html = f'<div class="todos">TODOs: {todo_list}</div>'

                out.write(f"""
<div class="file-block" id="file-{j}">
  <div class="file-header">
    <h3>{filename}</h3>
    <div class="file-meta">
      <span>{ext.upper().replace('.','')}</span>
      <span>{stats['total']:,} lines</span>
      <span>{stats['code']:,} code</span>
      <span>{stats['comment']:,} comments</span>
    </div>
    {classes_html}{funcs_html}{todos_html}
  </div>
  {highlighted}
</div>
""")

            out.write(f"""
  <div class="stats-footer">
    <strong>Document Stats</strong> &mdash;
    Files: {len(group)} &bull;
    Total lines: {total_lines:,}
  </div>
</div>
</body>
</html>
""")

        print(f"  ✓ {output_file}  ({len(group)} files)")

    print(f"\n  Created {num_docs} HTML documents in: {output_dir}")


# ─── Output: DOCX ─────────────────────────────────────────────────────────────

def create_docx_documents(consolidated_dir, output_dir, file_exts, num_docs=10):
    """Create .docx output documents (requires python-docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  python-docx not installed — falling back to text output.")
        print("  Install with: pip install python-docx\n")
        create_text_documents(consolidated_dir, output_dir, file_exts, num_docs)
        return

    print(f"\n[2/3] Creating DOCX documents...")

    files = get_all_files_by_extension(consolidated_dir, file_exts)
    total_files = len(files)

    if total_files == 0:
        print("  No files found!")
        return

    print(f"  Processing {total_files} files across {num_docs} documents")
    os.makedirs(output_dir, exist_ok=True)

    ext_str = "_".join([ext.replace(".", "") for ext in file_exts])
    doc_groups = split_into_documents(files, num_docs)

    for i, group in enumerate(doc_groups, 1):
        doc = Document()

        title_p = doc.add_heading(f"Code Files Collection — Part {i} of {num_docs}", 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info = doc.add_paragraph(
            f"File Types: {', '.join(file_exts)}     |     "
            f"Total Files: {len(group)}     |     "
            f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # TOC
        doc.add_heading("Table of Contents", level=1)
        for j, fp in enumerate(group, 1):
            doc.add_paragraph(f"{j}. {os.path.basename(fp)}")

        doc.add_page_break()

        for file_path in group:
            filename = os.path.basename(file_path)
            content = read_file_content(file_path)
            if not content:
                continue

            classes, functions = extract_symbols(file_path, content)
            todos = extract_todos(content)
            stats = file_stats(content)

            doc.add_heading(f"File: {filename}", level=2)
            doc.add_paragraph(
                f"Lines: {stats['total']:,} total  |  {stats['code']:,} code  |  "
                f"{stats['blank']:,} blank  |  {stats['comment']:,} comment"
            )
            doc.add_paragraph(f"Path: {file_path}")
            if classes:
                doc.add_paragraph(f"Classes: {', '.join(classes)}")
            if functions:
                doc.add_paragraph(f"Functions: {', '.join(functions)}")
            if todos:
                doc.add_paragraph(f"TODOs: {len(todos)} found")

            # Code content
            para = doc.add_paragraph(content)
            para.runs[0].font.name = "Courier New"
            para.runs[0].font.size = Pt(8)

            doc.add_paragraph("─" * 60)

        output_file = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.docx")
        doc.save(output_file)
        print(f"  ✓ {output_file}  ({len(group)} files)")

    print(f"\n  Created {num_docs} DOCX documents in: {output_dir}")


# ─── Index Document ────────────────────────────────────────────────────────────

def create_index_document(consolidated_dir, output_dir, file_exts, num_docs):
    """Create a _index.txt listing every file, its size, and which part it's in."""
    files = get_all_files_by_extension(consolidated_dir, file_exts)
    doc_groups = split_into_documents(files, num_docs)

    # Map file → part number
    file_to_part = {}
    for i, group in enumerate(doc_groups, 1):
        for f in group:
            file_to_part[f] = i

    index_path = os.path.join(output_dir, "_index.txt")
    total_size = 0
    total_lines = 0

    with open(index_path, "w", encoding="utf-8") as out:
        out.write(f"{'#'*80}\n")
        out.write(f"CODE FILES INDEX\n")
        out.write(f"Generated : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        out.write(f"Total Files: {len(files)}\n")
        out.write(f"{'#'*80}\n\n")
        out.write(f"{'PART':<6} {'SIZE':>8} {'LINES':>7}   FILENAME\n")
        out.write("─" * 70 + "\n")

        for f in files:
            part = file_to_part.get(f, "?")
            try:
                size = os.path.getsize(f)
                content = read_file_content(f)
                lines = len(content.split("\n")) if content else 0
            except Exception:
                size = 0
                lines = 0

            total_size += size
            total_lines += lines
            size_str = f"{size/1024:.1f}KB" if size >= 1024 else f"{size}B"
            out.write(f"{part:<6} {size_str:>8} {lines:>7}   {os.path.basename(f)}\n")

        out.write("\n" + "─" * 70 + "\n")
        total_size_str = f"{total_size/1024/1024:.2f}MB" if total_size >= 1024*1024 else f"{total_size/1024:.1f}KB"
        out.write(f"TOTAL: {len(files)} files  |  {total_size_str}  |  {total_lines:,} lines\n")

    print(f"  ✓ Index created: {index_path}")


# ─── Utilities ─────────────────────────────────────────────────────────────────

def read_file_content(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  Error reading {file_path}: {e}")
        return None


def split_into_documents(files, num_docs):
    if not files:
        return []
    num_docs = min(num_docs, len(files))  # Can't have more docs than files
    files_per_doc = len(files) // num_docs
    remainder = len(files) % num_docs
    docs, start = [], 0
    for i in range(num_docs):
        end = start + files_per_doc + (1 if i < remainder else 0)
        docs.append(files[start:end])
        start = end
    return docs


def parse_extensions(ext_input):
    ext_input = ext_input.replace(",", " ")
    normalized = []
    for ext in ext_input.split():
        ext = ext.strip().lower()
        if not ext.startswith("."):
            ext = "." + ext
        if ext and ext != ".":
            normalized.append(ext)
    return normalized


def parse_exclude_dirs(exclude_input):
    """Parse comma/space-separated list of folder names to exclude."""
    raw = exclude_input.replace(",", " ").split()
    return set(raw) if raw else set()


def print_supported_extensions():
    print("\nCommon supported extensions:")
    for lang, exts in COMMON_EXTENSIONS.items():
        print(f"  {lang:<25} {', '.join(exts)}")


# ─── Interactive Prompts ───────────────────────────────────────────────────────

def prompt_with_default(prompt, default):
    """Show a prompt with a default value and return user input or default."""
    if default:
        val = input(f"{prompt} [{default}]: ").strip()
        return val if val else default
    return input(f"{prompt}: ").strip()


def interactive_mode(saved_config):
    """Run the interactive prompt flow, pre-filling from saved config."""
    cfg = {}
    print("\n" + "─"*60)

    if saved_config:
        use_saved = input("  Found saved config. Use it? (Y/n): ").strip().lower()
        if use_saved != "n":
            print("  Loaded previous settings.\n")
            return saved_config

    # Project directory
    cfg["project_dir"] = prompt_with_default(
        "Project directory", saved_config.get("project_dir", ""))
    if not os.path.isdir(cfg["project_dir"]):
        print(f"  Error: '{cfg['project_dir']}' does not exist.")
        sys.exit(1)

    # Extensions
    print_supported_extensions()
    ext_input = prompt_with_default(
        "\nFile extensions (e.g. java, js, tsx)",
        " ".join(saved_config.get("extensions", [])))
    cfg["extensions"] = parse_extensions(ext_input)
    if not cfg["extensions"]:
        print("  Error: No valid extensions.")
        sys.exit(1)

    # Exclude dirs
    default_excl = ", ".join(sorted(DEFAULT_EXCLUDES))
    print(f"\n  Default excluded folders: {default_excl}")
    excl_input = prompt_with_default(
        "Additional folders to exclude (or press Enter to keep defaults)",
        ", ".join(saved_config.get("extra_excludes", [])))
    extra = parse_exclude_dirs(excl_input)
    cfg["extra_excludes"] = sorted(extra)
    cfg["exclude_dirs"] = sorted(DEFAULT_EXCLUDES | extra)

    # Consolidated folder
    ext_folder = "_".join([e.replace(".", "") for e in cfg["extensions"]])
    cfg["consolidated_dir"] = prompt_with_default(
        "Consolidated folder name",
        saved_config.get("consolidated_dir", f"consolidated_{ext_folder}"))

    # Output dir
    cfg["output_dir"] = prompt_with_default(
        "Output directory",
        saved_config.get("output_dir", "output_documents"))

    # Format
    print("\n  Output format:")
    print("    1. Text (.txt)")
    print("    2. DOCX (.docx)")
    print("    3. HTML with syntax highlighting (.html)")
    fmt = prompt_with_default("Choice", saved_config.get("format", "1"))
    cfg["format"] = fmt

    # Number of documents
    cfg["num_docs"] = int(prompt_with_default(
        "Number of output documents",
        str(saved_config.get("num_docs", 10))))

    # Index document
    make_index = prompt_with_default("Generate index document?", "y")
    cfg["make_index"] = make_index.lower() not in ("n", "no")

    return cfg


# ─── Main ──────────────────────────────────────────────────────────────────────

def build_arg_parser():
    parser = argparse.ArgumentParser(
        prog="code_to_documents",
        description="Code Files to Documents Converter v2.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Interactive mode (default)
  python code_to_documents.py

  # Fully CLI — process a Java project
  python code_to_documents.py /my/project --ext java --docs 5 --out java_docs

  # Multiple extensions, exclude extra folders, HTML output
  python code_to_documents.py /my/app --ext js jsx ts tsx --exclude dist coverage --format html

  # Dry run: see what would be processed without writing files
  python code_to_documents.py /my/project --ext py --dry-run
        """
    )
    parser.add_argument("project_dir", nargs="?", help="Path to source project directory")
    parser.add_argument("--ext", nargs="+", metavar="EXT",
                        help="File extensions to include (e.g. --ext java js jsx)")
    parser.add_argument("--exclude", nargs="*", metavar="DIR",
                        help="Extra folder names to exclude (adds to built-in defaults)")
    parser.add_argument("--consolidated", metavar="DIR",
                        help="Name of intermediate folder for copied files")
    parser.add_argument("--out", metavar="DIR", default="output_documents",
                        help="Output directory for documents (default: output_documents)")
    parser.add_argument("--format", choices=["txt", "docx", "html"], default="txt",
                        help="Output format: txt, docx, or html (default: txt)")
    parser.add_argument("--docs", type=int, default=10,
                        help="Number of output documents (default: 10)")
    parser.add_argument("--no-index", action="store_true",
                        help="Skip generating the _index.txt file")
    parser.add_argument("--dry-run", action="store_true",
                        help="Preview what would be processed without writing any files")
    parser.add_argument("--reset-config", action="store_true",
                        help="Ignore saved config and start fresh")
    return parser


def main():
    print("=" * 60)
    print("  Code Files to Documents Converter  v2.0")
    print("=" * 60)

    parser = build_arg_parser()
    args = parser.parse_args()

    # ── Determine config source ───────────────────────────────
    saved_config = {} if args.reset_config else load_config()

    if args.project_dir and args.ext:
        # CLI mode: all required args provided
        extra_excludes = set(args.exclude) if args.exclude else set()
        exclude_dirs = DEFAULT_EXCLUDES | extra_excludes
        ext_folder = "_".join([e.lstrip(".") for e in args.ext])

        cfg = {
            "project_dir": args.project_dir,
            "extensions": parse_extensions(" ".join(args.ext)),
            "exclude_dirs": sorted(exclude_dirs),
            "extra_excludes": sorted(extra_excludes),
            "consolidated_dir": args.consolidated or f"consolidated_{ext_folder}",
            "output_dir": args.out,
            "format": {"txt": "1", "docx": "2", "html": "3"}.get(args.format, "1"),
            "num_docs": args.docs,
            "make_index": not args.no_index,
        }
    else:
        # Interactive mode
        cfg = interactive_mode(saved_config)

    exclude_dirs = set(cfg.get("exclude_dirs", DEFAULT_EXCLUDES))

    # ── Dry run ───────────────────────────────────────────────
    if args.dry_run:
        count_files_preview(cfg["project_dir"], cfg["extensions"], exclude_dirs)
        return

    # ── Save config for next time ─────────────────────────────
    save_config(cfg)

    # ── Processing ────────────────────────────────────────────
    print("\n" + "=" * 50)
    print("  Processing...")
    print("=" * 50)

    # Step 1: Consolidate
    consolidate_files(
        cfg["project_dir"],
        cfg["consolidated_dir"],
        cfg["extensions"],
        exclude_dirs
    )

    # Step 2: Create documents
    fmt = cfg.get("format", "1")
    if fmt == "2":
        create_docx_documents(cfg["consolidated_dir"], cfg["output_dir"],
                               cfg["extensions"], cfg["num_docs"])
    elif fmt == "3":
        create_html_documents(cfg["consolidated_dir"], cfg["output_dir"],
                               cfg["extensions"], cfg["num_docs"])
    else:
        create_text_documents(cfg["consolidated_dir"], cfg["output_dir"],
                               cfg["extensions"], cfg["num_docs"])

    # Step 3: Index
    if cfg.get("make_index", True):
        print(f"\n[3/3] Creating index document...")
        create_index_document(cfg["consolidated_dir"], cfg["output_dir"],
                               cfg["extensions"], cfg["num_docs"])
    else:
        print(f"\n[3/3] Skipping index (--no-index)")

    # ── Summary ───────────────────────────────────────────────
    print("\n" + "=" * 50)
    print("  ALL DONE!")
    print("=" * 50)
    print(f"\n  Extensions       : {', '.join(cfg['extensions'])}")
    print(f"  Consolidated     : {cfg['consolidated_dir']}")
    print(f"  Output directory : {cfg['output_dir']}")
    print(f"  Documents created: {cfg['num_docs']}")
    print(f"  Format           : {{'1':'TXT','2':'DOCX','3':'HTML'}.get(cfg['format'],'TXT')}")
    print()


if __name__ == "__main__":
    main()
