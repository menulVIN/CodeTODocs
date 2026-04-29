#!/usr/bin/env python3
"""
Code Files to Documents Converter — GUI App v2.0 - By Venul Minsara
Double-click to run on Windows. No extra GUI dependencies needed (uses built-in tkinter).

Optional packages for extra features:
  pip install python-docx pygments tqdm
"""

import os
import sys
import shutil
import json
import re
import datetime
import threading
import queue
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# ── Optional dependencies ──────────────────────────────────────────────────────
try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False

try:
    from pygments import highlight
    from pygments.lexers import get_lexer_for_filename, TextLexer
    from pygments.formatters import HtmlFormatter
    from pygments.util import ClassNotFound
    HAS_PYGMENTS = True
except ImportError:
    HAS_PYGMENTS = False

# Optional modern ttk theme (ttkbootstrap). If installed, the app will prefer a
# dark bootstrap theme for a modern look. This is optional and not required.
try:
    import ttkbootstrap as tb
    HAS_TTKBOOTSTRAP = True
except Exception:
    HAS_TTKBOOTSTRAP = False

# ── Constants ──────────────────────────────────────────────────────────────────
CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "code_to_docs_config.json")

DEFAULT_EXCLUDES = {
    "node_modules", ".git", "__pycache__", "build", "dist",
    ".venv", "venv", "env", "target", ".idea", ".vscode",
    "out", "bin", "obj", ".next", ".nuxt", "coverage",
    ".gradle", ".mvn", "vendor", "bower_components",
}

EXTENSION_GROUPS = {
    "Java / Kotlin":      [".java", ".kt", ".scala"],
    "JavaScript":         [".js", ".jsx", ".mjs"],
    "TypeScript":         [".ts", ".tsx"],
    "Python":             [".py"],
    "C / C++":            [".c", ".cpp", ".h", ".hpp"],
    "C#":                 [".cs"],
    "Go":                 [".go"],
    "Rust":               [".rs"],
    "PHP / Ruby":         [".php", ".rb"],
    "Swift / ObjC":       [".swift", ".m", ".mm"],
    "Web (HTML/CSS)":     [".html", ".htm", ".css", ".scss", ".sass"],
    "Config / Data":      [".json", ".xml", ".yaml", ".yml"],
    "Docs (MD/TXT)":      [".md", ".txt"],
}

EXTRACTION_PATTERNS = {
    ".java":  [r"(?:public|private|protected|static|\s)+(?:class|interface|enum)\s+(\w+)",
               r"(?:public|private|protected|static|\s)+\w+\s+(\w+)\s*\("],
    ".kt":    [r"(?:class|object|interface|fun)\s+(\w+)"],
    ".py":    [r"^(?:class)\s+(\w+)", r"^(?:def)\s+(\w+)"],
    ".js":    [r"(?:function\s+(\w+)|class\s+(\w+)|const\s+(\w+)\s*=\s*(?:async\s*)?\()"],
    ".jsx":   [r"(?:function\s+(\w+)|class\s+(\w+)|const\s+(\w+)\s*=\s*(?:async\s*)?\()"],
    ".ts":    [r"(?:function\s+(\w+)|class\s+(\w+)|interface\s+(\w+)|const\s+(\w+)\s*=)"],
    ".tsx":   [r"(?:function\s+(\w+)|class\s+(\w+)|interface\s+(\w+)|const\s+(\w+)\s*=)"],
    ".cs":    [r"(?:class|interface|struct|enum)\s+(\w+)"],
    ".go":    [r"^func\s+(?:\(\w+\s+\*?\w+\)\s+)?(\w+)", r"^type\s+(\w+)\s+(?:struct|interface)"],
    ".rs":    [r"(?:fn|struct|enum|trait|impl)\s+(\w+)"],
    ".php":   [r"(?:function|class)\s+(\w+)"],
    ".rb":    [r"(?:def|class|module)\s+(\w+)"],
    ".swift": [r"(?:func|class|struct|enum|protocol)\s+(\w+)"],
}

# ── Colours & Fonts (theme) ────────────────────────────────────────────────────
BG        = "#0f1117"
BG2       = "#1a1d27"
BG3       = "#22263a"
PANEL     = "#1e2235"
ACCENT    = "#4f8ef7"
ACCENT2   = "#7c3aed"
SUCCESS   = "#22c55e"
WARNING   = "#f59e0b"
ERROR     = "#ef4444"
TEXT      = "#e2e8f0"
TEXT_DIM  = "#64748b"
WHITE     = "#ffffff"
BORDER    = "#2e3650"

FONT_UI   = ("Segoe UI", 11)
FONT_BOLD = ("Segoe UI", 11, "bold")
FONT_MONO = ("Consolas", 10)
FONT_H1   = ("Segoe UI", 16, "bold")
FONT_H2   = ("Segoe UI", 12, "bold")


# ══════════════════════════════════════════════════════════════════════════════
#  Core Processing Logic
# ══════════════════════════════════════════════════════════════════════════════

def read_file_content(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return None


def parse_extensions(ext_input):
    ext_input = ext_input.replace(",", " ")
    result = []
    for ext in ext_input.split():
        ext = ext.strip().lower()
        if not ext.startswith("."):
            ext = "." + ext
        if ext and ext != ".":
            result.append(ext)
    return result


def get_files(search_dir, extensions, exclude_dirs):
    normalized = []
    for ext in extensions:
        ext = ext.strip()
        if not ext.startswith("."):
            ext = "." + ext
        normalized.append(ext.lower())

    files = []
    for root, dirs, filenames in os.walk(search_dir):
        dirs[:] = [d for d in dirs if d not in exclude_dirs and not d.startswith(".")]
        for filename in filenames:
            if os.path.splitext(filename)[1].lower() in normalized:
                files.append(os.path.join(root, filename))
    return sorted(files)


def split_docs(files, num_docs):
    if not files:
        return []
    num_docs = min(num_docs, len(files))
    per = len(files) // num_docs
    rem = len(files) % num_docs
    groups, start = [], 0
    for i in range(num_docs):
        end = start + per + (1 if i < rem else 0)
        groups.append(files[start:end])
        start = end
    return groups


def extract_symbols(file_path, content):
    ext = os.path.splitext(file_path)[1].lower()
    patterns = EXTRACTION_PATTERNS.get(ext, [])
    if not patterns:
        return [], []
    SKIP = {"if","for","while","return","new","try","catch","import","from",
            "int","void","bool","string","var","let","const","self","cls"}
    classes, functions = [], []
    for line in content.split("\n"):
        for pat in patterns:
            for m in re.finditer(pat, line):
                for name in (g for g in m.groups() if g):
                    if len(name) > 1 and name not in SKIP:
                        if re.match(r"^[A-Z]", name) and name not in classes:
                            classes.append(name)
                        elif name not in functions:
                            functions.append(name)
    return classes[:8], functions[:12]


def extract_todos(content):
    todos = []
    for i, line in enumerate(content.split("\n"), 1):
        if re.search(r"\b(TODO|FIXME|HACK|XXX)\b", line, re.IGNORECASE):
            stripped = line.strip()
            todos.append((i, stripped[:80] + ("..." if len(stripped) > 80 else "")))
    return todos[:4]


def file_stats(content):
    lines = content.split("\n")
    total = len(lines)
    blank = sum(1 for l in lines if not l.strip())
    comment = sum(1 for l in lines if re.match(r"^\s*(//|#|/\*|\*|<!--)", l.strip()))
    return {"total": total, "code": total - blank - comment, "blank": blank, "comment": comment}


def format_file_block(file_path, content):
    filename = os.path.basename(file_path)
    file_type = os.path.splitext(filename)[1].upper().replace(".", "")
    classes, functions = extract_symbols(file_path, content)
    todos = extract_todos(content)
    stats = file_stats(content)

    out = f"{'='*80}\n"
    out += f"FILE: {filename}\n"
    out += f"TYPE: {file_type}   |   Lines: {stats['total']} total / {stats['code']} code / {stats['blank']} blank / {stats['comment']} comment\n"
    out += f"PATH: {file_path}\n"
    if classes:    out += f"CLASSES  : {', '.join(classes)}\n"
    if functions:  out += f"FUNCTIONS: {', '.join(functions)}\n"
    if todos:
        out += f"TODOs ({len(todos)}):\n"
        for ln, t in todos:
            out += f"  Line {ln}: {t}\n"
    out += f"{'='*80}\n\n{content}\n\n{'='*80}\nEND OF FILE: {filename}\n{'='*80}\n\n"
    return out, stats


def run_conversion(cfg, log_fn, progress_fn):
    """
    Core conversion logic. Runs in a background thread.
    log_fn(msg, tag)  — sends a message to the GUI log
    progress_fn(pct)  — updates progress bar 0-100
    """
    project_dir     = cfg["project_dir"]
    extensions      = cfg["extensions"]
    exclude_dirs    = cfg["exclude_dirs"]
    consolidated    = cfg["consolidated_dir"]
    output_dir      = cfg["output_dir"]
    fmt             = cfg["format"]          # "txt" | "docx" | "html"
    num_docs        = cfg["num_docs"]
    make_index      = cfg["make_index"]
    dry_run         = cfg.get("dry_run", False)

    log_fn("─" * 60, "dim")
    log_fn(f"  Project   : {project_dir}", "info")
    log_fn(f"  Extensions: {', '.join(extensions)}", "info")
    log_fn(f"  Exclude   : {', '.join(sorted(exclude_dirs))}", "dim")
    log_fn(f"  Output    : {output_dir}", "info")
    log_fn(f"  Format    : {fmt.upper()}   Docs: {num_docs}", "info")
    log_fn("─" * 60, "dim")

    # ── Discover files ──────────────────────────────────────
    log_fn("\n[1/3] Scanning project...", "header")
    all_files = get_files(project_dir, extensions, exclude_dirs)

    if not all_files:
        log_fn(f"  ✗ No files found matching: {extensions}", "error")
        log_fn("  Check your extensions or project directory.", "dim")
        progress_fn(-1)
        return

    log_fn(f"  ✓ Found {len(all_files)} files", "success")

    # Count by extension
    ext_counts = {}
    for f in all_files:
        e = os.path.splitext(f)[1].lower()
        ext_counts[e] = ext_counts.get(e, 0) + 1
    for e, c in sorted(ext_counts.items()):
        log_fn(f"    {e:<10} {c} files", "dim")

    if dry_run:
        log_fn("\n  DRY RUN — no files written.", "warning")
        progress_fn(100)
        return

    progress_fn(5)

    # ── Consolidate ─────────────────────────────────────────
    log_fn(f"\n[2/3] Consolidating → {consolidated}", "header")
    os.makedirs(consolidated, exist_ok=True)

    filename_counts = {}
    consolidated_files = []
    total = len(all_files)

    for idx, original_path in enumerate(all_files):
        filename = os.path.basename(original_path)
        if filename in filename_counts:
            filename_counts[filename] += 1
            rel = os.path.relpath(original_path, project_dir)
            parts = rel.split(os.sep)
            new_fn = f"{parts[-2]}_{filename}" if len(parts) > 1 else f"file_{filename_counts[filename]}_{filename}"
        else:
            filename_counts[filename] = 0
            new_fn = filename

        dest = os.path.join(consolidated, new_fn)
        counter = 1
        base, ext = os.path.splitext(new_fn)
        while os.path.exists(dest):
            dest = os.path.join(consolidated, f"{base}_{counter}{ext}")
            counter += 1

        shutil.copy2(original_path, dest)
        consolidated_files.append(dest)
        progress_fn(5 + int(35 * (idx + 1) / total))

    log_fn(f"  ✓ Copied {len(consolidated_files)} files", "success")

    # ── Create documents ─────────────────────────────────────
    log_fn(f"\n[3/3] Creating {fmt.upper()} documents...", "header")
    os.makedirs(output_dir, exist_ok=True)

    ready_files = get_files(consolidated, extensions, set())
    doc_groups  = split_docs(ready_files, num_docs)
    ext_str     = "_".join([e.replace(".", "") for e in extensions])
    total_docs  = len(doc_groups)
    file_to_part = {}
    for i, g in enumerate(doc_groups, 1):
        for f in g:
            file_to_part[f] = i

    doc_start_pct = 40

    if fmt == "html":
        _create_html(doc_groups, ext_str, output_dir, extensions, log_fn,
                     progress_fn, doc_start_pct)
    elif fmt == "docx":
        _create_docx(doc_groups, ext_str, output_dir, extensions, log_fn,
                     progress_fn, doc_start_pct)
    else:
        _create_txt(doc_groups, ext_str, output_dir, extensions, log_fn,
                    progress_fn, doc_start_pct)

    # ── Index ────────────────────────────────────────────────
    if make_index:
        log_fn("\n  Generating index...", "header")
        index_path = os.path.join(output_dir, "_index.txt")
        total_size, total_lines = 0, 0
        with open(index_path, "w", encoding="utf-8") as idx_f:
            idx_f.write(f"{'#'*80}\nCODE FILES INDEX\n")
            idx_f.write(f"Generated : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            idx_f.write(f"Total Files: {len(ready_files)}\n{'#'*80}\n\n")
            idx_f.write(f"{'PART':<6} {'SIZE':>8} {'LINES':>7}   FILENAME\n")
            idx_f.write("─" * 70 + "\n")
            for f in ready_files:
                part = file_to_part.get(f, "?")
                try:
                    size = os.path.getsize(f)
                    content = read_file_content(f)
                    lines = len(content.split("\n")) if content else 0
                except Exception:
                    size = lines = 0
                total_size += size
                total_lines += lines
                size_str = f"{size/1024:.1f}KB" if size >= 1024 else f"{size}B"
                idx_f.write(f"{part:<6} {size_str:>8} {lines:>7}   {os.path.basename(f)}\n")
            idx_f.write("\n" + "─"*70 + "\n")
            mb = total_size/1024/1024
            sz = f"{mb:.2f}MB" if mb >= 1 else f"{total_size/1024:.1f}KB"
            idx_f.write(f"TOTAL: {len(ready_files)} files  |  {sz}  |  {total_lines:,} lines\n")
        log_fn(f"  ✓ Index: {index_path}", "success")

    progress_fn(100)
    log_fn("\n" + "═"*60, "dim")
    log_fn("  ✓  ALL DONE!", "success")
    log_fn(f"  Output folder: {output_dir}", "info")
    log_fn("═"*60, "dim")


def _create_txt(doc_groups, ext_str, output_dir, extensions, log_fn, progress_fn, base_pct):
    total = len(doc_groups)
    for i, group in enumerate(doc_groups, 1):
        path = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.txt")
        doc_stats = {"total": 0, "code": 0, "blank": 0, "comment": 0}
        with open(path, "w", encoding="utf-8") as out:
            out.write(f"{'#'*80}\nCode Files Collection — Part {i} of {total}\n")
            out.write(f"File Types : {', '.join(extensions)}\n")
            out.write(f"Files      : {len(group)}\n")
            out.write(f"Generated  : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n{'#'*80}\n\n")
            out.write("TABLE OF CONTENTS\n" + "─"*40 + "\n")
            for j, fp in enumerate(group, 1):
                out.write(f"  {j:>3}. {os.path.basename(fp)}\n")
            out.write("\n" + "─"*40 + "\n\n")
            for fp in group:
                content = read_file_content(fp)
                if content:
                    formatted, stats = format_file_block(fp, content)
                    out.write(formatted)
                    for k in doc_stats:
                        doc_stats[k] += stats.get(k, 0)
            out.write(f"\n{'#'*80}\nDOCUMENT STATS\n")
            out.write(f"  Total    : {doc_stats['total']:,}\n  Code     : {doc_stats['code']:,}\n")
            out.write(f"  Blank    : {doc_stats['blank']:,}\n  Comments : {doc_stats['comment']:,}\n")
            out.write(f"{'#'*80}\n")
        log_fn(f"  ✓ {os.path.basename(path)}  ({len(group)} files)", "success")
        progress_fn(base_pct + int(55 * i / total))


def _create_docx(doc_groups, ext_str, output_dir, extensions, log_fn, progress_fn, base_pct):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log_fn("  ✗ python-docx not installed — falling back to TXT", "warning")
        log_fn("    Run: pip install python-docx", "dim")
        _create_txt(doc_groups, ext_str, output_dir, extensions, log_fn, progress_fn, base_pct)
        return

    total = len(doc_groups)
    for i, group in enumerate(doc_groups, 1):
        doc = Document()
        p = doc.add_heading(f"Code Files Collection — Part {i} of {total}", 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Types: {', '.join(extensions)}   |   Files: {len(group)}   |   {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_heading("Table of Contents", level=1)
        for j, fp in enumerate(group, 1):
            doc.add_paragraph(f"{j}. {os.path.basename(fp)}")
        doc.add_page_break()
        for fp in group:
            filename = os.path.basename(fp)
            content = read_file_content(fp)
            if not content:
                continue
            classes, functions = extract_symbols(fp, content)
            stats = file_stats(content)
            doc.add_heading(f"File: {filename}", level=2)
            doc.add_paragraph(f"Lines: {stats['total']:,} total  |  {stats['code']:,} code  |  {stats['blank']:,} blank")
            if classes:   doc.add_paragraph(f"Classes  : {', '.join(classes)}")
            if functions: doc.add_paragraph(f"Functions: {', '.join(functions)}")
            para = doc.add_paragraph(content)
            para.runs[0].font.name = "Courier New"
            para.runs[0].font.size = Pt(8)
            doc.add_paragraph("─" * 60)
        path = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.docx")
        doc.save(path)
        log_fn(f"  ✓ {os.path.basename(path)}  ({len(group)} files)", "success")
        progress_fn(base_pct + int(55 * i / total))


def _create_html(doc_groups, ext_str, output_dir, extensions, log_fn, progress_fn, base_pct):
    if not HAS_PYGMENTS:
        log_fn("  ✗ pygments not installed — falling back to TXT", "warning")
        log_fn("    Run: pip install pygments", "dim")
        _create_txt(doc_groups, ext_str, output_dir, extensions, log_fn, progress_fn, base_pct)
        return

    formatter = HtmlFormatter(style="monokai", linenos=True, cssclass="code")
    css = formatter.get_style_defs(".code")
    total = len(doc_groups)

    for i, group in enumerate(doc_groups, 1):
        toc = "".join(f'<li><a href="#f{j}">{os.path.basename(fp)}</a></li>'
                      for j, fp in enumerate(group, 1))
        path = os.path.join(output_dir, f"{ext_str}_part_{i:02d}.html")
        with open(path, "w", encoding="utf-8") as out:
            out.write(f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>Code — Part {i} of {total}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',sans-serif;background:#0f1117;color:#e2e8f0}}
.sidebar{{position:fixed;top:0;left:0;width:260px;height:100vh;background:#1a1d27;
  overflow-y:auto;padding:16px;border-right:1px solid #2e3650}}
.sidebar h2{{color:#4f8ef7;font-size:12px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:12px}}
.sidebar ol{{padding-left:16px}}
.sidebar li{{margin-bottom:4px}}
.sidebar a{{color:#94a3b8;text-decoration:none;font-size:11px;word-break:break-all}}
.sidebar a:hover{{color:#e2e8f0}}
.content{{margin-left:276px;padding:24px;max-width:1200px}}
h1{{color:#4f8ef7;font-size:18px;margin-bottom:4px}}
.meta{{color:#64748b;font-size:12px;margin-bottom:28px}}
.file-block{{margin-bottom:32px;border:1px solid #2e3650;border-radius:8px;overflow:hidden}}
.file-header{{background:#1e2235;padding:10px 16px;border-bottom:1px solid #2e3650}}
.file-header h3{{color:#7dd3fc;font-size:13px;margin-bottom:4px}}
.file-meta{{font-size:11px;color:#64748b}}
.file-meta span{{margin-right:14px}}
.symbols{{font-size:11px;color:#86efac;margin-top:3px}}
.todos{{font-size:11px;color:#fca5a5;margin-top:3px}}
.highlight{{margin:0;border-radius:0}}
.stats{{background:#1a1d27;padding:14px 20px;border-radius:8px;
  font-size:12px;color:#64748b;margin-top:28px}}
{css}
</style></head><body>
<div class="sidebar">
  <h2>Part {i} / {total}</h2>
  <p style="color:#64748b;font-size:11px;margin-bottom:12px">{len(group)} files</p>
  <ol>{toc}</ol>
</div>
<div class="content">
  <h1>Code Collection — Part {i} of {total}</h1>
  <div class="meta">{', '.join(extensions)} &bull; {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
""")
            total_lines = 0
            for j, fp in enumerate(group, 1):
                content = read_file_content(fp)
                if not content:
                    continue
                filename = os.path.basename(fp)
                classes, functions = extract_symbols(fp, content)
                todos = extract_todos(content)
                stats = file_stats(content)
                total_lines += stats["total"]
                try:
                    lexer = get_lexer_for_filename(filename)
                except ClassNotFound:
                    lexer = TextLexer()
                hi = highlight(content, lexer, formatter)
                sym = f'<div class="symbols">Classes: {", ".join(classes)}</div>' if classes else ""
                sym += f'<div class="symbols">Functions: {", ".join(functions)}</div>' if functions else ""
                td = (f'<div class="todos">TODOs: ' +
                      " &bull; ".join(f"L{ln}: {t[:50]}" for ln, t in todos[:2]) +
                      "</div>") if todos else ""
                out.write(f"""<div class="file-block" id="f{j}">
<div class="file-header">
  <h3>{filename}</h3>
  <div class="file-meta">
    <span>{os.path.splitext(filename)[1].upper().replace('.','')}</span>
    <span>{stats['total']:,} lines</span>
    <span>{stats['code']:,} code</span>
    <span>{stats['comment']:,} comments</span>
  </div>{sym}{td}
</div>{hi}</div>\n""")
            out.write(f'<div class="stats">Files: {len(group)} &bull; Total lines: {total_lines:,}</div>')
            out.write("</div></body></html>")
        log_fn(f"  ✓ {os.path.basename(path)}  ({len(group)} files)", "success")
        progress_fn(base_pct + int(55 * i / total))


# ══════════════════════════════════════════════════════════════════════════════
#  Config helpers
# ══════════════════════════════════════════════════════════════════════════════

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE) as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_config(cfg):
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(cfg, f, indent=2)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
#  GUI Application
# ══════════════════════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Code Files to Documents Converter  v2.0")
        self.geometry("1100x760")
        self.minsize(900, 640)
        self.configure(bg=BG)
        self._log_queue = queue.Queue()
        self._running   = False

        self._setup_styles()
        self._build_ui()
        self._load_saved()
        self.after(100, self._drain_log_queue)

    # ── Styles ─────────────────────────────────────────────────────────────────
    def _setup_styles(self):
        # Prefer ttkbootstrap dark theme when available for a modern appearance.
        if HAS_TTKBOOTSTRAP:
            try:
                s = tb.Style(theme="cyborg")
            except Exception:
                s = ttk.Style(self)
                s.theme_use("clam")
        else:
            s = ttk.Style(self)
            s.theme_use("clam")

        s.configure(".", background=BG, foreground=TEXT, font=FONT_UI,
                    borderwidth=0, relief="flat")
        s.configure("TFrame",     background=BG)
        s.configure("Panel.TFrame", background=PANEL)
        s.configure("TLabel",     background=BG, foreground=TEXT)
        s.configure("Panel.TLabel", background=PANEL, foreground=TEXT)
        s.configure("Dim.TLabel", background=PANEL, foreground=TEXT_DIM, font=("Segoe UI", 9))
        s.configure("H1.TLabel",  background=BG,    foreground=TEXT, font=FONT_H1)
        s.configure("H2.TLabel",  background=PANEL, foreground=ACCENT, font=FONT_H2)

        s.configure("TEntry",
                    fieldbackground=BG3, foreground=TEXT,
                    insertcolor=TEXT, borderwidth=1, relief="flat")
        s.map("TEntry", fieldbackground=[("focus", BG3)])

        s.configure("TCheckbutton",
                    background=PANEL, foreground=TEXT,
                    indicatorcolor=BG3, indicatorrelief="flat")
        s.map("TCheckbutton",
              indicatorcolor=[("selected", ACCENT)],
              foreground=[("active", ACCENT)])

        s.configure("TRadiobutton",
                    background=PANEL, foreground=TEXT)
        s.map("TRadiobutton",
              foreground=[("active", ACCENT), ("selected", ACCENT)])

        s.configure("Accent.TButton",
                  background=ACCENT, foreground=WHITE,
                  font=("Segoe UI", 11, "bold"),
                  padding=(18, 8), relief="flat", borderwidth=0)
        s.map("Accent.TButton",
              background=[("active", "#6ba3ff"), ("disabled", "#334155")],
              foreground=[("disabled", TEXT_DIM)])

        s.configure("Small.TButton",
                    background=BG3, foreground=TEXT,
                    font=FONT_UI, padding=(8, 4), relief="flat")
        s.map("Small.TButton",
              background=[("active", BORDER)])

        s.configure("TProgressbar",
                    background=ACCENT, troughcolor=BG3,
                    borderwidth=0, thickness=6)

        # Slim, themed vertical scrollbar
        s.configure("Vertical.TScrollbar",
                troughcolor=BG3, background=BG3, arrowcolor=TEXT_DIM)

        s.configure("TSpinbox",
                    fieldbackground=BG3, foreground=TEXT,
                    insertcolor=TEXT, arrowcolor=TEXT_DIM,
                    borderwidth=1, relief="flat")

        s.configure("TNotebook",            background=BG,    borderwidth=0)
        s.configure("TNotebook.Tab",        background=BG3,   foreground=TEXT_DIM,
                    padding=(12, 6), font=FONT_UI)
        s.map("TNotebook.Tab",
              background=[("selected", PANEL)],
              foreground=[("selected", ACCENT)])

    # ── UI Layout ──────────────────────────────────────────────────────────────
    def _build_ui(self):
        # ── Header bar ──────────────────────────────────
        header = tk.Frame(self, bg=BG2, height=54)
        header.pack(fill="x")
        header.pack_propagate(False)

        # thin separator below header for visual separation
        sep = tk.Frame(self, bg=BORDER, height=1)
        sep.pack(fill="x")

        tk.Label(header, text="⬡", font=("Segoe UI", 20), bg=BG2,
                 fg=ACCENT).pack(side="left", padx=(18, 8), pady=8)
        tk.Label(header, text="Code Files to Documents", font=("Segoe UI", 14, "bold"),
                 bg=BG2, fg=TEXT).pack(side="left", pady=8)
        tk.Label(header, text="v2.0", font=("Segoe UI", 10),
                 bg=BG2, fg=TEXT_DIM).pack(side="left", padx=6, pady=8)

        # Optional deps status
        deps = []
        if HAS_PYGMENTS: deps.append("pygments ✓")
        if HAS_TQDM:     deps.append("tqdm ✓")
        try:
            import docx
            deps.append("python-docx ✓")
        except ImportError:
            pass
        if deps:
            tk.Label(header, text="  |  " + "   ".join(deps),
                     font=("Segoe UI", 9), bg=BG2, fg=SUCCESS).pack(side="left", pady=8)

        # ── Main body: left panel + right log ───────────
        body = ttk.Frame(self)
        body.pack(fill="both", expand=True, padx=0, pady=0)

        # Left scrollable panel
        left_outer = tk.Frame(body, bg=BG, width=480)
        left_outer.pack(side="left", fill="y")
        left_outer.pack_propagate(False)

        canvas = tk.Canvas(left_outer, bg=BG, highlightthickness=0, width=480)
        scrollbar = ttk.Scrollbar(left_outer, orient="vertical", command=canvas.yview, style="Vertical.TScrollbar")
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        self._left = ttk.Frame(canvas, style="TFrame")
        canvas_win = canvas.create_window((0, 0), window=self._left, anchor="nw")

        def _on_frame_configure(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def _on_canvas_configure(e):
            canvas.itemconfig(canvas_win, width=e.width)
        self._left.bind("<Configure>", _on_frame_configure)
        canvas.bind("<Configure>", _on_canvas_configure)
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(-1*(e.delta//120), "units"))

        # Right log panel
        right = tk.Frame(body, bg=BG)
        right.pack(side="left", fill="both", expand=True)

        self._build_left_panel(self._left)
        self._build_right_panel(right)

        # ── Bottom bar ──────────────────────────────────
        bottom = tk.Frame(self, bg=BG2, height=64)
        bottom.pack(fill="x", side="bottom")
        bottom.pack_propagate(False)

        self._progress = ttk.Progressbar(bottom, mode="determinate", maximum=100)
        self._progress.pack(fill="x", padx=20, pady=(10, 0))

        btn_frame = tk.Frame(bottom, bg=BG2)
        btn_frame.pack(fill="x", padx=20, pady=(4, 0))

        self._status_var = tk.StringVar(value="Ready")
        tk.Label(btn_frame, textvariable=self._status_var,
                 font=("Segoe UI", 9), bg=BG2, fg=TEXT_DIM).pack(side="left")

        # Use a compact emoji icon for a modern affordance; the label keeps spacing
        # for readability across platforms.
        self._run_btn = ttk.Button(btn_frame, text="⚡  Run Conversion",
                       style="Accent.TButton", command=self._on_run)
        self._run_btn.pack(side="right", padx=(8, 0))

        ttk.Button(btn_frame, text="🧭  Dry Run",
               style="Small.TButton", command=self._on_dry_run).pack(side="right")

    def _section(self, parent, title):
        """Create a labelled section card."""
        # card with subtle border and extra inner padding for a modern 'card' feel
        card = tk.Frame(parent, bg=PANEL, bd=0, highlightbackground=BORDER, highlightthickness=1)
        card.pack(fill="x", padx=12, pady=(8, 0))
        tk.Label(card, text=title, font=FONT_H2, bg=PANEL, fg=ACCENT).pack(
            anchor="w", padx=16, pady=(12, 8))
        inner = tk.Frame(card, bg=PANEL)
        inner.pack(fill="x", padx=16, pady=(0, 14))
        return inner

    def _row(self, parent, label, widget_factory):
        """Label + widget side by side."""
        row = tk.Frame(parent, bg=PANEL)
        row.pack(fill="x", pady=3)
        tk.Label(row, text=label, font=FONT_UI, bg=PANEL, fg=TEXT,
                 width=16, anchor="w").pack(side="left")
        w = widget_factory(row)
        w.pack(side="left", fill="x", expand=True)
        return w

    def _browse_row(self, parent, label, var, folder=True):
        row = tk.Frame(parent, bg=PANEL)
        row.pack(fill="x", pady=3)
        tk.Label(row, text=label, font=FONT_UI, bg=PANEL, fg=TEXT,
                 width=16, anchor="w").pack(side="left")
        entry = ttk.Entry(row, textvariable=var)
        entry.pack(side="left", fill="x", expand=True)

        def browse():
            if folder:
                path = filedialog.askdirectory(title=f"Select {label}")
            else:
                path = filedialog.asksaveasfilename(title=f"Save {label}")
            if path:
                var.set(path)

        ttk.Button(row, text="Browse…", style="Small.TButton",
                   command=browse).pack(side="left", padx=(6, 0))

    def _build_left_panel(self, parent):
        tk.Label(parent, text="Configuration", font=FONT_H1,
                 bg=BG, fg=TEXT).pack(anchor="w", padx=14, pady=(14, 6))

        # ── Directories ─────────────────────────────────
        sec = self._section(parent, "Directories")
        self._project_var = tk.StringVar()
        self._output_var  = tk.StringVar(value="output_documents")
        self._consol_var  = tk.StringVar()
        self._browse_row(sec, "Project Dir", self._project_var)
        self._browse_row(sec, "Output Dir",  self._output_var)
        self._browse_row(sec, "Consolidated", self._consol_var)
        tk.Label(sec, text="Leave 'Consolidated' blank to auto-name",
                 font=("Segoe UI", 8), bg=PANEL, fg=TEXT_DIM).pack(anchor="w", pady=(2, 0))

        # ── File Extensions ──────────────────────────────
        sec2 = self._section(parent, "File Extensions")
        tk.Label(sec2, text="Tick groups or type custom extensions below:",
                 font=("Segoe UI", 9), bg=PANEL, fg=TEXT_DIM).pack(anchor="w", pady=(0, 6))

        self._ext_vars = {}
        grid = tk.Frame(sec2, bg=PANEL)
        grid.pack(fill="x")
        col, row_n = 0, 0
        for i, (group, _) in enumerate(EXTENSION_GROUPS.items()):
            var = tk.BooleanVar()
            self._ext_vars[group] = var
            cb = ttk.Checkbutton(grid, text=group, variable=var, style="TCheckbutton")
            cb.grid(row=row_n, column=col, sticky="w", padx=(0, 12), pady=2)
            col += 1
            if col >= 2:
                col = 0
                row_n += 1

        tk.Label(sec2, text="Custom extensions (comma separated):",
                 font=("Segoe UI", 9), bg=PANEL, fg=TEXT_DIM).pack(anchor="w", pady=(8, 2))
        self._custom_ext = tk.StringVar()
        ttk.Entry(sec2, textvariable=self._custom_ext).pack(fill="x")
        tk.Label(sec2, text="e.g.   dart, sol, vue, svelte",
                 font=("Segoe UI", 8), bg=PANEL, fg=TEXT_DIM).pack(anchor="w")

        # ── Exclude Folders ──────────────────────────────
        sec3 = self._section(parent, "Exclude Folders")
        tk.Label(sec3,
                 text="Default excluded: " + ", ".join(sorted(DEFAULT_EXCLUDES)[:8]) + " ...",
                 font=("Segoe UI", 8), bg=PANEL, fg=TEXT_DIM,
                 wraplength=400, justify="left").pack(anchor="w", pady=(0, 6))
        tk.Label(sec3, text="Additional folders to exclude:",
                 font=("Segoe UI", 9), bg=PANEL, fg=TEXT).pack(anchor="w", pady=(0, 2))
        self._exclude_var = tk.StringVar()
        ttk.Entry(sec3, textvariable=self._exclude_var).pack(fill="x")
        tk.Label(sec3, text="Folder names separated by commas   e.g.  tests, fixtures, docs",
                 font=("Segoe UI", 8), bg=PANEL, fg=TEXT_DIM).pack(anchor="w")

        # ── Output Format ────────────────────────────────
        sec4 = self._section(parent, "Output Format")
        self._fmt_var = tk.StringVar(value="txt")
        fmts = [
            ("txt",  "Plain Text (.txt)          — no extra packages"),
            ("docx", "Microsoft Word (.docx)     — needs python-docx"),
            ("html", "HTML Syntax Highlighted    — needs pygments"),
        ]
        for val, label in fmts:
            ttk.Radiobutton(sec4, text=label, variable=self._fmt_var,
                            value=val).pack(anchor="w", pady=2)

        # ── Options ──────────────────────────────────────
        sec5 = self._section(parent, "Options")
        row1 = tk.Frame(sec5, bg=PANEL)
        row1.pack(fill="x", pady=3)
        tk.Label(row1, text="Number of docs:", font=FONT_UI, bg=PANEL, fg=TEXT,
                 width=16, anchor="w").pack(side="left")
        self._numdocs_var = tk.IntVar(value=10)
        ttk.Spinbox(row1, from_=1, to=100, textvariable=self._numdocs_var,
                    width=6, font=FONT_UI).pack(side="left")

        self._index_var   = tk.BooleanVar(value=True)
        self._open_var    = tk.BooleanVar(value=True)

        ttk.Checkbutton(sec5, text="Generate _index.txt  (lists every file, size, lines, part)",
                        variable=self._index_var).pack(anchor="w", pady=2)
        ttk.Checkbutton(sec5, text="Open output folder when done",
                        variable=self._open_var).pack(anchor="w", pady=2)

        tk.Frame(parent, bg=BG, height=16).pack()

    def _build_right_panel(self, parent):
        header = tk.Frame(parent, bg=BG2, height=36)
        header.pack(fill="x")
        header.pack_propagate(False)
        tk.Label(header, text="  Log Output", font=FONT_BOLD,
                 bg=BG2, fg=TEXT).pack(side="left", padx=8, pady=6)
        ttk.Button(header, text="Clear", style="Small.TButton",
                   command=self._clear_log).pack(side="right", padx=8, pady=5)

        self._log = scrolledtext.ScrolledText(
            parent, font=FONT_MONO, bg="#0a0c14", fg="#c9d1d9",
            insertbackground=TEXT, wrap="word", state="disabled",
            relief="flat", borderwidth=0, padx=12, pady=10)
        self._log.pack(fill="both", expand=True)

        # Tag colours
        self._log.tag_config("header",  foreground="#7dd3fc", font=("Consolas", 9, "bold"))
        self._log.tag_config("success", foreground=SUCCESS)
        self._log.tag_config("warning", foreground=WARNING)
        self._log.tag_config("error",   foreground=ERROR)
        self._log.tag_config("info",    foreground="#94a3b8")
        self._log.tag_config("dim",     foreground=TEXT_DIM)
        self._log.tag_config("normal",  foreground="#c9d1d9")

    # ── Helpers ────────────────────────────────────────────────────────────────
    def _log_append(self, msg, tag="normal"):
        self._log.configure(state="normal")
        self._log.insert("end", msg + "\n", tag)
        self._log.see("end")
        self._log.configure(state="disabled")

    def _clear_log(self):
        self._log.configure(state="normal")
        self._log.delete("1.0", "end")
        self._log.configure(state="disabled")

    def _drain_log_queue(self):
        try:
            while True:
                msg, tag = self._log_queue.get_nowait()
                if msg == "__DONE__":
                    self._on_done(tag == "ok")
                elif msg == "__PROGRESS__":
                    val = float(tag)
                    if val < 0:
                        self._progress["value"] = 0
                    else:
                        self._progress["value"] = val
                else:
                    self._log_append(msg, tag)
        except queue.Empty:
            pass
        self.after(80, self._drain_log_queue)

    def _load_saved(self):
        cfg = load_config()
        if not cfg:
            return
        if cfg.get("project_dir"):     self._project_var.set(cfg["project_dir"])
        if cfg.get("output_dir"):      self._output_var.set(cfg["output_dir"])
        if cfg.get("consolidated_dir"):self._consol_var.set(cfg["consolidated_dir"])
        if cfg.get("format"):          self._fmt_var.set(cfg["format"])
        if cfg.get("num_docs"):        self._numdocs_var.set(cfg["num_docs"])
        if "make_index" in cfg:        self._index_var.set(cfg["make_index"])
        if cfg.get("custom_ext"):      self._custom_ext.set(cfg["custom_ext"])
        if cfg.get("extra_excludes"):
            self._exclude_var.set(", ".join(cfg["extra_excludes"]))
        # Restore checked groups
        for group, exts in EXTENSION_GROUPS.items():
            saved_exts = cfg.get("extensions", [])
            if any(e in saved_exts for e in exts):
                self._ext_vars[group].set(True)

    def _collect_extensions(self):
        exts = []
        for group, group_exts in EXTENSION_GROUPS.items():
            if self._ext_vars[group].get():
                exts.extend(group_exts)
        custom = self._custom_ext.get().strip()
        if custom:
            exts.extend(parse_extensions(custom))
        return list(dict.fromkeys(exts))  # deduplicate, preserve order

    def _collect_config(self, dry_run=False):
        project_dir = self._project_var.get().strip()
        if not project_dir:
            messagebox.showerror("Missing", "Please select a Project Directory.")
            return None
        if not os.path.isdir(project_dir):
            messagebox.showerror("Not Found", f"Directory not found:\n{project_dir}")
            return None

        extensions = self._collect_extensions()
        if not extensions:
            messagebox.showerror("Missing", "Please select at least one file extension.")
            return None

        extra_exc = set()
        for e in self._exclude_var.get().replace(",", " ").split():
            e = e.strip()
            if e:
                extra_exc.add(e)

        ext_folder = "_".join(e.lstrip(".") for e in extensions[:4])
        consol = self._consol_var.get().strip() or f"consolidated_{ext_folder}"

        cfg = {
            "project_dir":      project_dir,
            "extensions":       extensions,
            "exclude_dirs":     sorted(DEFAULT_EXCLUDES | extra_exc),
            "extra_excludes":   sorted(extra_exc),
            "consolidated_dir": consol,
            "output_dir":       self._output_var.get().strip() or "output_documents",
            "format":           self._fmt_var.get(),
            "num_docs":         self._numdocs_var.get(),
            "make_index":       self._index_var.get(),
            "custom_ext":       self._custom_ext.get().strip(),
            "dry_run":          dry_run,
        }
        return cfg

    # ── Actions ────────────────────────────────────────────────────────────────
    def _on_run(self):
        if self._running:
            return
        cfg = self._collect_config(dry_run=False)
        if cfg is None:
            return
        save_config(cfg)
        self._start_thread(cfg)

    def _on_dry_run(self):
        if self._running:
            return
        cfg = self._collect_config(dry_run=True)
        if cfg is None:
            return
        self._start_thread(cfg)

    def _start_thread(self, cfg):
        self._running = True
        self._run_btn.configure(state="disabled")
        self._progress["value"] = 0
        self._status_var.set("Running…")
        self._clear_log()

        def log_fn(msg, tag="normal"):
            self._log_queue.put((msg, tag))

        def progress_fn(pct):
            self._log_queue.put(("__PROGRESS__", str(pct)))

        def worker():
            try:
                run_conversion(cfg, log_fn, progress_fn)
                self._log_queue.put(("__DONE__", "ok"))
            except Exception as ex:
                log_fn(f"\n  ✗ Error: {ex}", "error")
                import traceback
                log_fn(traceback.format_exc(), "dim")
                self._log_queue.put(("__DONE__", "err"))

        threading.Thread(target=worker, daemon=True).start()

    def _on_done(self, success):
        self._running = False
        self._run_btn.configure(state="normal")
        if success:
            self._status_var.set("Done ✓")
            if self._open_var.get():
                out = self._output_var.get().strip() or "output_documents"
                if os.path.isdir(out):
                    os.startfile(out)
        else:
            self._status_var.set("Finished with errors")


# ══════════════════════════════════════════════════════════════════════════════
#  Entry point
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # On Windows, hide the console window when double-clicking
    if sys.platform == "win32":
        try:
            import ctypes
            ctypes.windll.user32.ShowWindow(
                ctypes.windll.kernel32.GetConsoleWindow(), 0)
        except Exception:
            pass

    app = App()
    app.mainloop()
